"""
Format Checker API
==================

API endpoint for checking document formatting against Vietnamese
government standards via RAG.

POST /format-check/
    Upload a .docx file and get formatting analysis report
"""

from __future__ import annotations

import io
import logging
import uuid
from pathlib import Path
from typing import Optional

from fastapi import APIRouter, Depends, HTTPException, status, UploadFile, File, Form
from fastapi.responses import StreamingResponse
from sqlalchemy.ext.asyncio import AsyncSession
from pydantic import BaseModel

from app.core.deps import get_db, get_current_active_user
from app.models.user import User
from app.models.chat_file import ChatFile
from app.models.format_metadata import FormatMetadata
from app.services.storage_service import get_storage_service

logger = logging.getLogger(__name__)

router = APIRouter(prefix="/format-check", tags=["format-check"])


class FormatCheckResponse(BaseModel):
    report: str
    file_name: str
    issues_count: int
    chat_file_id: Optional[str] = None
    error: Optional[str] = None


async def _get_accessible_workspaces(db: AsyncSession, user: User) -> list[uuid.UUID]:
    """Get all knowledge base IDs the user has access to."""
    import uuid
    from sqlalchemy import select, or_
    from app.models.knowledge_base import KnowledgeBase
    from app.models.tenant import TenantUser

    if user.is_superadmin:
        result = await db.execute(select(KnowledgeBase.id))
        return list(result.scalars().all())

    tenant_result = await db.execute(
        select(TenantUser.tenant_id).where(TenantUser.user_id == user.id)
    )
    user_tenant_ids = list(tenant_result.scalars().all())

    query = select(KnowledgeBase.id).where(
        or_(
            KnowledgeBase.visibility == "public",
            KnowledgeBase.owner_id == user.id,
            KnowledgeBase.tenant_id.in_(user_tenant_ids) if user_tenant_ids else False,
        )
    )
    result = await db.execute(query)
    return list(result.scalars().all())


def _docx_to_markdown(file_path: str) -> str:
    """Convert docx content to markdown text."""
    from docx import Document
    from docx.shared import Pt

    doc = Document(file_path)
    markdown_lines = []

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            markdown_lines.append("")
            continue

        style_name = para.style.name if para.style else "Normal"

        if style_name.startswith("Heading"):
            level = style_name[-1] if style_name[-1].isdigit() else "1"
            markdown_lines.append(f"{'#' * int(level)} {text}")
        elif style_name == "Title":
            markdown_lines.append(f"# {text}")
        elif style_name == "Subtitle":
            markdown_lines.append(f"## {text}")
        else:
            markdown_lines.append(text)

    for table in doc.tables:
        markdown_lines.append("")
        markdown_lines.append(
            "| "
            + " | ".join([cell.text.strip()[:50] for cell in table.rows[0].cells])
            + " |"
        )
        markdown_lines.append(
            "| " + " | ".join(["---" for _ in table.rows[0].cells]) + " |"
        )
        for row in table.rows[1:]:
            markdown_lines.append(
                "| " + " | ".join([cell.text.strip()[:50] for cell in row.cells]) + " |"
            )
        markdown_lines.append("")

    return "\n".join(markdown_lines)


@router.post("/", response_model=FormatCheckResponse)
async def check_document_format(
    file: UploadFile = File(...),
    session_id: str = Form(...),
    db: AsyncSession = Depends(get_db),
    user: User = Depends(get_current_active_user),
):
    """
    Check document formatting against Vietnamese government standards.

    - Accepts .docx files only
    - Extracts formatting metadata (margins, line spacing, fonts)
    - Compares against standards via RAG
    - Saves file and metadata to MinIO + DB
    - Returns detailed formatting report with chat_file_id
    """
    ext = Path(file.filename).suffix.lower()
    if ext != ".docx":
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="Only .docx files are supported for format checking",
        )

    try:
        content = await file.read()
    except (ConnectionResetError, OSError) as exc:
        logger.warning(f"Client disconnected during upload: {exc}")
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="Upload interrupted — client disconnected. Please retry.",
        )

    if len(content) > 10 * 1024 * 1024:  # 10MB limit
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="File too large. Max size: 10MB",
        )

    import tempfile
    import os

    with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp:
        tmp.write(content)
        tmp_path = tmp.name

    try:
        workspace_ids = await _get_accessible_workspaces(db, user)
        if not workspace_ids:
            raise HTTPException(
                status_code=status.HTTP_403_FORBIDDEN,
                detail="You don't have access to any workspace",
            )

        workspace_id = workspace_ids[0]

        from app.services.agents.docx_formatter_tools import (
            extract_docx_format,
            analyze_format_issues,
        )

        extracted = await extract_docx_format(tmp_path)

        if extracted.get("error"):
            return FormatCheckResponse(
                report=f"Không thể đọc tệp {file.filename}. Vui lòng đảm bảo đây là tệp Word (.docx) hợp lệ.",
                file_name=file.filename,
                issues_count=0,
                error=extracted["error"],
            )

        issues = analyze_format_issues(extracted)

        margins = extracted.get("margins", {})
        font_samples = extracted.get("font_samples", [])

        from collections import Counter

        font_sizes = [f["font_size"] for f in font_samples if f.get("font_size")]
        size_counts = Counter(font_sizes)
        most_common_sizes = size_counts.most_common(3)

        report_lines = [
            f"## BÁO CÁO KIỂM TRA ĐỊNH DẠNG: {file.filename}",
            "",
            "### THÔNG TIN ĐỊNH DẠNG:",
            f"- Căn lề: Trên {margins.get('top', 'N/A')}cm, Dưới {margins.get('bottom', 'N/A')}cm, Trái {margins.get('left', 'N/A')}cm, Phải {margins.get('right', 'N/A')}cm",
            f"- Cỡ chữ phổ biến: {', '.join([f'{s}pt' for s, _ in most_common_sizes]) if most_common_sizes else 'N/A'}",
            f"- Số đoạn: {extracted.get('paragraph_count', 'N/A')}",
            f"- Số bảng: {extracted.get('table_count', 'N/A')}",
            "",
            f"### VẤN ĐỀ PHÁT HIỆN: {len(issues)} vấn đề",
        ]

        if issues:
            for issue in issues:
                report_lines.append(
                    f"- [{issue['severity'].upper()}] {issue['detail']}"
                )
                report_lines.append(f"  → {issue['suggestion']}")
        else:
            report_lines.append("Không phát hiện vấn đề định dạng nghiêm trọng.")

        report_lines.extend(
            [
                "",
                "### GHI CHÚ:",
                "- Chuẩn căn lề: Trên 2cm, Dưới 2cm, Trái 3cm, Phải 2cm",
                "- Chuẩn cỡ chữ: 13pt cho nội dung, 14pt cho tiêu đề",
                "- Chuẩn khoảng cách dòng: 1.5 dòng",
            ]
        )

        report = "\n".join(report_lines)

        storage = get_storage_service()
        file_uuid = uuid.uuid4()

        minio_original_key = f"kb_{workspace_id}/chat_file_{file_uuid}.docx"
        minio_markdown_key = f"kb_{workspace_id}/chat_file_{file_uuid}.md"

        await storage.upload_file(
            minio_original_key,
            content,
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )

        markdown_content = _docx_to_markdown(tmp_path)
        await storage.upload_markdown(workspace_id, file_uuid, markdown_content)

        chat_file = ChatFile(
            session_id=uuid.UUID(session_id),
            user_id=user.id,
            workspace_id=workspace_id,
            file_name=f"chat_file_{file_uuid}.docx",
            original_filename=file.filename or "unknown.docx",
            file_type="docx",
            file_size=len(content),
            minio_original_key=minio_original_key,
            minio_markdown_key=minio_markdown_key,
            markdown_content=markdown_content,
            report=report,
            issues_count=len(issues),
        )
        db.add(chat_file)
        await db.flush()

        format_metadata = FormatMetadata(
            chat_file_id=chat_file.id,
            format_data=extracted,
        )
        db.add(format_metadata)
        await db.commit()

        return FormatCheckResponse(
            report=report,
            file_name=file.filename,
            issues_count=len(issues),
            chat_file_id=str(chat_file.id),
            error=None,
        )

    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"[check_document_format] Failed: {e}")
        await db.rollback()
        return FormatCheckResponse(
            report="Đã xảy ra lỗi khi kiểm tra định dạng. Vui lòng thử lại.",
            file_name=file.filename,
            issues_count=0,
            error=str(e),
        )
    finally:
        try:
            os.unlink(tmp_path)
        except Exception:
            pass


@router.post("/stream/")
async def check_document_format_stream(
    file: UploadFile = File(...),
    session_id: str = Form(...),
    db: AsyncSession = Depends(get_db),
    user: User = Depends(get_current_active_user),
):
    """
    Check document formatting with streaming response.

    Returns SSE stream for real-time progress updates.
    """
    ext = Path(file.filename).suffix.lower()
    if ext != ".docx":
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="Only .docx files are supported for format checking",
        )

    try:
        content = await file.read()
    except (ConnectionResetError, OSError) as exc:
        logger.warning(f"Client disconnected during upload: {exc}")
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="Upload interrupted — client disconnected. Please retry.",
        )

    if len(content) > 10 * 1024 * 1024:  # 10MB limit
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="File too large. Max size: 10MB",
        )

    import tempfile
    import os

    with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp:
        tmp.write(content)
        tmp_path = tmp.name

    async def event_generator():
        try:
            from app.services.agents.docx_formatter_tools import (
                extract_docx_format,
                analyze_format_issues,
            )

            yield f'data: {{"step": "extracting", "message": "Đang trích xuất thông tin định dạng..."}}\n\n'

            extracted = await extract_docx_format(tmp_path)

            if extracted.get("error"):
                yield f'data: {{"step": "error", "message": "Không thể đọc tệp"}}\n\n'
                return

            yield f'data: {{"step": "analyzing", "message": "Đang phân tích vấn đề..."}}\n\n'

            issues = analyze_format_issues(extracted)

            yield f'data: {{"step": "complete", "message": "Hoàn thành", "issues_count": {len(issues)}}}\n\n'

        except Exception as e:
            logger.error(f"[check_document_format_stream] Error: {e}")
            yield f'data: {{"step": "error", "message": "{str(e)}"}}\n\n'
        finally:
            try:
                os.unlink(tmp_path)
            except Exception:
                pass

    return StreamingResponse(
        event_generator(),
        media_type="text/event-stream",
        headers={
            "Cache-Control": "no-cache",
            "Connection": "keep-alive",
        },
    )
