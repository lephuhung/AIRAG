"""
Docx Formatter Tools
====================

Tools for extracting formatting information from .docx files
and consulting RAG for Vietnamese government document standards.

Tools:
    - extract_docx_format: Extract formatting metadata from a .docx file
    - rag_lookup_format_standards: Look up formatting standards via RAG
    - compare_format_with_standards: Compare extracted format with standards
"""

from __future__ import annotations

import logging
from typing import Any

logger = logging.getLogger(__name__)

from docx import Document
from docx.shared import Inches, Pt, Cm, Emu
from docx.enum.text import WD_LINE_SPACING, WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import re


MARGIN_STANDARDS = {
    "top": 2.0,  # cm - thường dùng cho văn bản hành chính
    "bottom": 2.0,
    "left": 3.0,  # D CPO bên trái để đóng dấu
    "right": 2.0,
}

FONT_SIZE_STANDARDS = {
    "body": 13,  # 13pt cho nội dung chính (T CV 01/2014)
    "title": 14,  # 14pt cho tiêu đề
    "subtitle": 12,
    "small": 12,
}

LINE_SPACING_STANDARDS = {
    "exact": 20,  # pt - khoảng cách dòng cố định
    "multiple": 1.5,  # 1.5 dòng cho văn bản thông thường
    "double": 2.0,  # 2.0 dòng cho văn bản pháp luật
}


def cm_to_inches(cm: float) -> float:
    return cm / 2.54


def pt_to_cm(pt: float) -> float:
    return pt * 0.035277


def inches_to_cm(inches: float) -> float:
    return inches * 2.54


def extract_docx_format_sync(file_path: str) -> dict[str, Any]:
    """Synchronous wrapper for extract_docx_format — runs CPU-bound docx parsing in thread pool."""
    import asyncio
    return asyncio.run(extract_docx_format(file_path))


async def extract_docx_format(file_path: str) -> dict[str, Any]:
    """
    Extract comprehensive formatting information from a .docx file.

    Returns:
        dict with keys:
            - margins: dict with top, bottom, left, right margins (in cm)
            - line_spacing: dict with spacing_type, spacing_value (in pt or ratio)
            - font_samples: list of font info from first paragraphs
            - page_size: dict with width, height (in cm)
            - paragraph_count: total paragraphs
            - table_count: total tables
            - has_header_footer: bool
            - styles_used: list of style names
    """
    try:
        doc = Document(file_path)

        # Extract page/section properties
        section = doc.sections[0]

        # Margins (convert to cm)
        margins = {
            "top": round(inches_to_cm(section.top_margin.inches), 2),
            "bottom": round(inches_to_cm(section.bottom_margin.inches), 2),
            "left": round(inches_to_cm(section.left_margin.inches), 2),
            "right": round(inches_to_cm(section.right_margin.inches), 2),
        }

        # Page size (convert to cm)
        page_size = {
            "width": round(inches_to_cm(section.page_width.inches), 2),
            "height": round(inches_to_cm(section.page_height.inches), 2),
        }

        # Line spacing analysis
        line_spacing_samples = []
        for i, para in enumerate(doc.paragraphs[:20]):
            pf = para.paragraph_format
            spacing_info = {
                "paragraph_index": i,
                "style": para.style.name if para.style else "Normal",
                "line_spacing_type": None,
                "line_spacing_value": None,
                "space_before": pt_to_cm(pf.space_before.pt) if pf.space_before else 0,
                "space_after": pt_to_cm(pf.space_after.pt) if pf.space_after else 0,
            }

            if pf.line_spacing_rule:
                if pf.line_spacing_rule == WD_LINE_SPACING.EXACTLY:
                    spacing_info["line_spacing_type"] = "exact"
                    spacing_info["line_spacing_value"] = (
                        pf.line_spacing.pt if pf.line_spacing else None
                    )
                elif pf.line_spacing_rule == WD_LINE_SPACING.MULTIPLE:
                    spacing_info["line_spacing_type"] = "multiple"
                    spacing_info["line_spacing_value"] = (
                        pf.line_spacing if pf.line_spacing else 1.0
                    )
                elif pf.line_spacing_rule == WD_LINE_SPACING.AT_LEAST:
                    spacing_info["line_spacing_type"] = "at_least"
                    spacing_info["line_spacing_value"] = (
                        pf.line_spacing.pt if pf.line_spacing else None
                    )
                elif pf.line_spacing_rule == WD_LINE_SPACING.AUTO:
                    spacing_info["line_spacing_type"] = "auto"
                    spacing_info["line_spacing_value"] = (
                        pf.line_spacing if pf.line_spacing else 1.0
                    )

            line_spacing_samples.append(spacing_info)

        # Font analysis - first check rPrDefault for document-wide defaults
        default_font = None
        default_font_size = None

        try:
            import zipfile
            with zipfile.ZipFile(file_path, 'r') as z:
                with z.open('word/styles.xml') as f:
                    styles_xml = f.read().decode('utf-8')
                    # Extract default font from rPrDefault
                    import re
                    rpr_match = re.search(r'<w:rPrDefault>.*?<w:rFonts[^>]*w:ascii="([^"]*)"', styles_xml, re.DOTALL)
                    if rpr_match:
                        default_font = rpr_match.group(1)
                    # Extract default font size
                    sz_match = re.search(r'<w:sz w:val="(\d+)"', styles_xml)
                    if sz_match:
                        default_font_size = int(sz_match.group(1)) / 2  # Convert half-points to pt
        except Exception as e:
            logger.warning(f"[extract_docx_format] Could not read styles.xml defaults: {e}")

        # Now analyze runs (may inherit defaults if not explicitly set)
        font_samples = []
        for i, para in enumerate(doc.paragraphs[:30]):
            for run in para.runs:
                if run.text.strip():
                    font_info = {
                        "paragraph_index": i,
                        "font_name": run.font.name or default_font,
                        "font_size": run.font.size.pt if run.font.size else default_font_size,
                        "bold": run.font.bold,
                        "italic": run.font.italic,
                        "underline": run.font.underline,
                        "text_preview": run.text[:50] + "..."
                        if len(run.text) > 50
                        else run.text,
                        "is_default": run.font.name is None,  # True if using document default
                    }
                    font_samples.append(font_info)
                    if len(font_samples) >= 50:
                        break
            if len(font_samples) >= 50:
                break

        # Collect unique styles
        styles_used = list(
            set(para.style.name for para in doc.paragraphs if para.style)
        )[:20]  # Limit to 20 styles

        # Check header/footer
        has_header = any(
            section.header.is_linked_to_previous == False for section in doc.sections
        )
        has_footer = any(
            section.footer.is_linked_to_previous == False for section in doc.sections
        )

        return {
            "margins": margins,
            "page_size": page_size,
            "line_spacing": line_spacing_samples,
            "font_samples": font_samples,
            "paragraph_count": len(doc.paragraphs),
            "table_count": len(doc.tables),
            "has_header": has_header,
            "has_footer": has_footer,
            "styles_used": styles_used,
            "error": None,
        }

    except Exception as e:
        logger.error(
            f"[extract_docx_format] Failed to extract format from {file_path}: {e}"
        )
        return {
            "margins": {},
            "page_size": {},
            "line_spacing": [],
            "font_samples": [],
            "paragraph_count": 0,
            "table_count": 0,
            "has_header": False,
            "has_footer": False,
            "styles_used": [],
            "error": str(e),
        }


def analyze_format_issues(extracted: dict) -> list[dict]:
    """
    Analyze extracted formatting and identify issues compared to Vietnamese standards.

    Returns list of issues, each with: type, severity, detail, suggestion
    """
    issues = []

    margins = extracted.get("margins", {})
    line_spacing = extracted.get("line_spacing", [])
    font_samples = extracted.get("font_samples", [])
    styles_used = extracted.get("styles_used", [])

    # Check margins
    for side, standard in MARGIN_STANDARDS.items():
        actual = margins.get(side)
        if actual is not None:
            diff = abs(actual - standard)
            if diff > 0.3:  # More than 3mm deviation
                issues.append(
                    {
                        "type": "margin",
                        "severity": "warning" if diff < 1.0 else "error",
                        "detail": f"Căn lề {side}: {actual}cm (chuẩn: {standard}cm)",
                        "suggestion": f"Điều chỉnh căn lề {side} về {standard}cm",
                    }
                )

    # Analyze line spacing
    spacing_types = {}
    for sp in line_spacing:
        if sp["line_spacing_type"]:
            stype = sp["line_spacing_type"]
            sval = sp["line_spacing_value"]
            key = f"{stype}_{sval}"
            spacing_types[key] = spacing_types.get(key, 0) + 1

    if spacing_types:
        most_common_spacing = max(spacing_types, key=spacing_types.get)
        if (
            "multiple_1.5" not in most_common_spacing
            and "multiple_1.0" not in most_common_spacing
        ):
            issues.append(
                {
                    "type": "line_spacing",
                    "severity": "warning",
                    "detail": f"Khoảng cách dòng phổ biến: {most_common_spacing}",
                    "suggestion": "Nên sử dụng khoảng cách 1.5 dòng cho văn bản hành chính",
                }
            )

    # Check font sizes
    font_sizes = [f["font_size"] for f in font_samples if f.get("font_size")]
    if font_sizes:
        body_fonts = [s for s in font_sizes if s and 11 <= s <= 14]
        non_standard = [s for s in body_fonts if s not in [12, 13]]

        if non_standard:
            most_common = (
                max(set(body_fonts), key=body_fonts.count) if body_fonts else None
            )
            issues.append(
                {
                    "type": "font_size",
                    "severity": "info",
                    "detail": f"Cỡ chữ phổ biến: {most_common}pt (chuẩn: 13pt)",
                    "suggestion": f"Cỡ chữ nên dùng: 13pt cho nội dung, 14pt cho tiêu đề",
                }
            )

    # Check for mixed fonts
    unique_fonts = set(f["font_name"] for f in font_samples if f.get("font_name"))
    if len(unique_fonts) > 3:
        issues.append(
            {
                "type": "font_mixed",
                "severity": "warning",
                "detail": f"Đang sử dụng {len(unique_fonts)} loại font: {', '.join(list(unique_fonts)[:5])}",
                "suggestion": "Nên chỉ sử dụng 1-2 loại font chính (Times New Roman, Arial)",
            }
        )

    return issues


async def rag_lookup_format_standards(
    query: str,
    workspace_ids: list[int],
    db: "AsyncSession",
    top_k: int = 5,
) -> list[str]:
    """
    Look up Vietnamese government document formatting standards via RAG.

    Uses DeepRetriever to search for relevant standards in the workspace.

    Returns list of relevant context strings.
    """
    from sqlalchemy import select
    from app.models.document import Document, DocumentStatus
    from app.services.retrieval.deep_retriever import DeepRetriever
    from app.services.embedding.vector_store import VectorStore
    from app.core.config import settings

    try:
        # Get indexed documents in workspace
        result = await db.execute(
            select(Document.id).where(
                Document.workspace_id.in_(workspace_ids),
                Document.status == DocumentStatus.INDEXED,
            )
        )
        doc_ids = [row[0] for row in result.fetchall()]

        if not doc_ids:
            return []

        # Use DeepRetriever for hybrid search
        retriever = DeepRetriever(
            workspace_ids=workspace_ids,
            vector_store=VectorStore(),
            top_k=top_k,
        )

        retrieval_result = await retriever.query(
            question=query,
            mode="hybrid",
            document_ids=doc_ids,
            include_images=False,
        )

        # Extract context from chunks
        contexts = []
        for chunk in retrieval_result.chunks[:top_k]:
            content = chunk.get("content", "")
            source = chunk.get("source_file", "")
            page = chunk.get("page_no", "")
            contexts.append(f"[{source} - Trang {page}]:\n{content}")

        return contexts

    except Exception as e:
        logger.error(f"[rag_lookup_format_standards] RAG lookup failed: {e}")
        return []


async def compare_format_with_standards(
    extracted: dict,
    workspace_ids: list[int],
    db: "AsyncSession",
) -> dict:
    """
    Compare extracted formatting with standards found via RAG.

    Returns a comprehensive comparison report.
    """
    issues = analyze_format_issues(extracted)

    # Look up relevant standards via RAG
    format_query = extracted.get(
        "query", "quy chuẩn trình bày văn bản hành chính Việt Nam"
    )
    standards_contexts = await rag_lookup_format_standards(
        query=format_query,
        workspace_ids=workspace_ids,
        db=db,
        top_k=5,
    )

    return {
        "extracted": extracted,
        "issues": issues,
        "standards_contexts": standards_contexts,
        "standards_count": len(standards_contexts),
    }
